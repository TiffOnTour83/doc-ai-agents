"""Generate sample documents for trying the pipeline.

Creates two revisions of a fictional operations manual so every feature —
extraction, validation, harmonization, taxonomy, and the revision diff —
has something real to chew on. All content is invented.

    python examples/make_samples.py
    python -m doclogic run examples/manual_rev_a.docx --profile aviation-manual
    python -m doclogic diff examples/manual_rev_a.docx examples/manual_rev_b.docx
"""

from pathlib import Path

import docx

HERE = Path(__file__).resolve().parent


def rev_a() -> None:
    d = docx.Document()
    d.add_paragraph("Working draft. Distribution pending.")  # orphan text
    d.add_heading("1 Program Approval", level=1)
    d.add_paragraph("The approval workflow covers submission and authorization.")
    d.add_heading("1.1 Submission Process", level=2)
    d.add_paragraph("Curriculum revisions are submitted for administrator review.")
    d.add_heading("2 Qualification Standards", level=1)
    d.add_paragraph("Duty position standards are defined per fleet and seat.")
    d.add_heading("2.1.1 Currency Requirements", level=3)  # level jump
    d.add_paragraph("Currency windows and proficiency requirements.")
    d.add_heading("3 Records", level=1)  # empty section
    d.add_heading("4 Training Curriculum", level=1)
    d.add_paragraph("Course outlines and task analysis methodology.")
    d.save(HERE / "manual_rev_a.docx")


def rev_b() -> None:
    d = docx.Document()
    d.add_heading("1 Program Approval", level=1)
    d.add_paragraph("The approval workflow covers submission and authorization.")
    d.add_heading("1.1 Submission Workflow", level=2)  # renamed
    d.add_paragraph("Curriculum revisions are submitted for administrator review.")
    d.add_heading("2 Training Curriculum", level=1)  # moved up, renumbered
    d.add_paragraph("Course outlines and task analysis methodology.")
    d.add_heading("3 Qualification Standards", level=1)  # renumbered
    d.add_paragraph("Duty position standards are defined per fleet and seat.")
    d.add_heading("4 Instructor Certification", level=1)  # added
    d.add_paragraph("Certification pathway for instructors and evaluators.")
    # "Records" removed
    d.save(HERE / "manual_rev_b.docx")


if __name__ == "__main__":
    rev_a()
    rev_b()
    print("wrote", HERE / "manual_rev_a.docx")
    print("wrote", HERE / "manual_rev_b.docx")
