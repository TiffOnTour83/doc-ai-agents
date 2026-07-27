"""End-to-end and unit tests for the Doc Logic pipeline."""

import json
import sys
from pathlib import Path

import docx
import pytest

sys.path.append(str(Path(__file__).resolve().parents[1]))

from doclogic.audit import classify, rtm_skeleton, structured_json  # noqa: E402
from doclogic.diff import diff  # noqa: E402
from doclogic.extract import extract  # noqa: E402
from doclogic.gate import blocking_errors, record_decision  # noqa: E402
from doclogic.harmonize import harmonize  # noqa: E402
from doclogic.profiles import load_profile  # noqa: E402
from doclogic.suggest import RuleBasedSuggester  # noqa: E402


def _write(tmp_path, name, build):
    d = docx.Document()
    build(d)
    path = tmp_path / name
    d.save(path)
    return str(path)


@pytest.fixture
def messy(tmp_path):
    def build(d):
        d.add_paragraph("Orphan preamble.")
        d.add_heading("1 Approval", level=1)
        d.add_paragraph("Approval workflow and administrator authorization.")
        d.add_heading("1.1.1 Deep Dive", level=3)  # jump
        d.add_paragraph("Too deep too fast.")
        d.add_heading("2 Records", level=1)  # empty later? has body below
        d.add_paragraph("Record retention and audit documentation.")
        d.add_heading("2 Records", level=1)  # duplicate
        d.add_paragraph("Duplicate content.")
    return _write(tmp_path, "messy.docx", build)


def test_extract_styles_and_numbers(messy):
    extraction = extract(messy)
    assert extraction.sections[0].number == "1"
    assert extraction.sections[0].title == "Approval"
    assert extraction.orphan_paragraphs == ["Orphan preamble."]
    assert len(extraction.sha256) == 64


def test_suggester_flags_all_issues(messy):
    proposal = RuleBasedSuggester().propose(extract(messy), load_profile("default"))
    checks = {i.check for i in proposal.issues}
    assert {"level-jump", "duplicate-title", "orphan-text"} <= checks
    assert blocking_errors(proposal)  # errors present -> gate blocks


def test_aviation_profile_requires_numbering(tmp_path):
    def build(d):
        d.add_heading("Unnumbered Heading", level=1)
        d.add_paragraph("Body.")
    path = _write(tmp_path, "styles_only.docx", build)
    proposal = RuleBasedSuggester().propose(
        extract(path), load_profile("aviation-manual")
    )
    assert any(i.check == "numbering-required" for i in proposal.issues)


def test_harmonizer_fixes_and_logs_everything(messy):
    harmonized = harmonize(extract(messy), load_profile("default"))
    levels = [s.level for s in harmonized.sections]
    for previous, current in zip([0] + levels, levels):
        assert current <= previous + 1  # no jumps remain
    actions = {c.action for c in harmonized.changes}
    assert "promote" in actions
    assert "renumber" in actions
    numbers = [s.number for s in harmonized.sections]
    assert numbers == sorted(numbers, key=lambda n: [int(x) for x in n.split(".")])


def test_taxonomy_classifies_with_evidence(messy):
    harmonized = harmonize(extract(messy), load_profile("default"))
    taxonomy = classify(harmonized, load_profile("default"))
    record_section = next(
        c for c in taxonomy["classifications"] if "Records" in c["title"]
    )
    assert record_section["category"] == "Records"
    assert record_section["evidence"]  # transparent: names the keywords


def test_rtm_and_json_artifacts(messy):
    harmonized = harmonize(extract(messy), load_profile("default"))
    csv_bytes = rtm_skeleton(harmonized, "messy.docx")
    assert b"section_number" in csv_bytes and b"Pending" in csv_bytes
    tree = json.loads(structured_json(harmonized, "messy.docx"))
    assert tree["source"] == "messy.docx"
    assert tree["sections"][0]["children"] is not None


def test_decision_is_recorded_as_artifact(messy, tmp_path):
    proposal = RuleBasedSuggester().propose(extract(messy), load_profile("default"))
    decision = record_decision(proposal, True, "T. Analyst", "ok", tmp_path)
    log = (tmp_path / "decisions.jsonl").read_text().strip().splitlines()
    assert len(log) == 1
    entry = json.loads(log[0])
    assert entry["approver"] == "T. Analyst"
    assert entry["proposal_digest"] == proposal.digest()


def test_diff_detects_all_change_kinds(tmp_path):
    def rev_a(d):
        d.add_heading("1 Approval", level=1)
        d.add_paragraph("a")
        d.add_heading("2 Submission Process", level=1)
        d.add_paragraph("b")
        d.add_heading("3 Records", level=1)
        d.add_paragraph("c")

    def rev_b(d):
        d.add_heading("1 Approval", level=1)
        d.add_paragraph("a")
        d.add_heading("2 Submission Workflow", level=1)  # renamed
        d.add_paragraph("b")
        d.add_heading("3 Instructor Certification", level=1)  # added
        d.add_paragraph("d")
        # Records removed

    old = extract(_write(tmp_path, "a.docx", rev_a))
    new = extract(_write(tmp_path, "b.docx", rev_b))
    kinds = {e.kind for e in diff(old, new)}
    assert {"renamed", "added", "removed"} <= kinds


def test_cli_gate_blocks_then_produces_artifacts(messy, tmp_path, capsys):
    from doclogic.__main__ import main

    # Without approval: stops at the gate.
    assert main(["run", messy, "--out", str(tmp_path)]) == 0
    assert "Gate: stopped" in capsys.readouterr().out

    # With approval but blocking errors: refuses without override.
    assert main(["run", messy, "--approve", "T", "--out", str(tmp_path)]) == 1

    # With override: full artifact set + manifest.
    assert main(
        ["run", messy, "--approve", "T", "--override", "--out", str(tmp_path)]
    ) == 0
    produced = {p.name for p in tmp_path.iterdir()}
    assert {
        "messy_structured.json", "messy_rtm_skeleton.csv",
        "messy_taxonomy.json", "messy_audit_report.txt",
        "messy_manifest.json", "decisions.jsonl",
    } <= produced
    manifest = json.loads((tmp_path / "messy_manifest.json").read_text())
    assert manifest["approval"]["approver"] == "T"
    assert manifest["source"]["sha256"]
