"""Tests for the format auditor and cross-document standardization."""

import sys
from pathlib import Path

import docx
from docx.shared import Pt

sys.path.append(str(Path(__file__).resolve().parents[1]))

from doclogic.formats import (  # noqa: E402
    audit_document,
    profile_document,
    standardize,
)


def _doc(tmp_path, name, body_font="Calibri", body_size=11, odd_paragraph=None):
    d = docx.Document()
    d.add_heading("1 Section One", level=1)
    for text in ["First body paragraph.", "Second body paragraph."]:
        p = d.add_paragraph()
        run = p.add_run(text)
        run.font.name = body_font
        run.font.size = Pt(body_size)
    if odd_paragraph:
        p = d.add_paragraph()
        run = p.add_run(odd_paragraph)
        run.font.name = "Comic Sans MS"
        run.font.size = Pt(14)
    path = tmp_path / name
    d.save(path)
    return str(path)


def test_within_document_deviation_is_named(tmp_path):
    path = _doc(tmp_path, "doc.docx", odd_paragraph="I am different.")
    findings = audit_document(profile_document(path))
    assert findings, "expected the Comic Sans paragraph to be flagged"
    assert any("Comic Sans MS" in f.detail for f in findings)
    assert all(f.kind == "body" for f in findings)


def test_consistent_document_is_clean(tmp_path):
    path = _doc(tmp_path, "clean.docx")
    assert audit_document(profile_document(path)) == []


def test_cross_document_standard_is_common_denominator(tmp_path):
    profiles = [
        profile_document(_doc(tmp_path, "a.docx", body_font="Calibri")),
        profile_document(_doc(tmp_path, "b.docx", body_font="Calibri")),
        profile_document(_doc(tmp_path, "c.docx", body_font="Arial")),
    ]
    standard, findings = standardize(profiles)
    assert standard["body"]["font"] == "Calibri"
    assert standard["body"]["documents_agreeing"] == 2
    assert any("c.docx" in f.detail and "Arial" in f.detail for f in findings)
