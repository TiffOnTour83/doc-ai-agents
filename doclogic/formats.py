"""Format auditor — consistency analysis for how documents *look*.

Two levels:

* Within one document: do headings at the same level share a font and size?
  Does body text drift between typefaces? Findings name the exact paragraphs.
* Across documents: fingerprint each document's dominant formatting, derive
  the common-denominator standard, and report every deviation from it.

The cross-document flow is the methodology in miniature: find the common
denominators, make them the standard, and list exactly what has drifted.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

import docx

from .extract import _HEADING_STYLE  # reuse the style-name detector


@dataclass
class ParaFormat:
    text: str
    kind: str            # "heading-1", "heading-2", ... or "body"
    font: str            # dominant run font, or "(inherited)"
    size: str            # dominant run size in pt, or "(inherited)"
    style: str


@dataclass
class FormatProfile:
    source: str
    paragraphs: list[ParaFormat]
    # kind -> most common (font, size) with counts
    dominant: dict[str, tuple[str, str]] = field(default_factory=dict)


@dataclass
class FormatFinding:
    scope: str    # document name, or "cross-document"
    kind: str     # "heading-1", "body", ...
    detail: str


def _run_font(paragraph) -> tuple[str, str]:
    """Dominant explicit run font/size, falling back to the paragraph style."""
    fonts = Counter()
    sizes = Counter()
    for run in paragraph.runs:
        if run.font.name:
            fonts[run.font.name] += len(run.text)
        if run.font.size is not None:
            sizes[str(run.font.size.pt)] += len(run.text)
    style_font = None
    style_size = None
    if paragraph.style is not None:
        style_font = paragraph.style.font.name
        if paragraph.style.font.size is not None:
            style_size = str(paragraph.style.font.size.pt)
    font = fonts.most_common(1)[0][0] if fonts else (style_font or "(inherited)")
    size = sizes.most_common(1)[0][0] if sizes else (style_size or "(inherited)")
    return font, size


def profile_document(path: str | Path) -> FormatProfile:
    path = Path(path)
    document = docx.Document(str(path))
    paragraphs: list[ParaFormat] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "") if paragraph.style else ""
        match = _HEADING_STYLE.search(style_name)
        kind = f"heading-{match.group(1)}" if match else "body"
        font, size = _run_font(paragraph)
        paragraphs.append(
            ParaFormat(text=text[:60], kind=kind, font=font, size=size, style=style_name)
        )

    profile = FormatProfile(source=path.name, paragraphs=paragraphs)
    by_kind: dict[str, Counter] = {}
    for para in paragraphs:
        by_kind.setdefault(para.kind, Counter())[(para.font, para.size)] += 1
    for kind, counter in by_kind.items():
        profile.dominant[kind] = counter.most_common(1)[0][0]
    return profile


def audit_document(profile: FormatProfile) -> list[FormatFinding]:
    """Within-document consistency: deviations from each kind's dominant look."""
    findings: list[FormatFinding] = []
    for kind, (dom_font, dom_size) in sorted(profile.dominant.items()):
        deviants = [
            p for p in profile.paragraphs
            if p.kind == kind and (p.font, p.size) != (dom_font, dom_size)
        ]
        for para in deviants:
            findings.append(
                FormatFinding(
                    scope=profile.source,
                    kind=kind,
                    detail=(
                        f"'{para.text}' uses {para.font}/{para.size}pt "
                        f"(dominant for {kind}: {dom_font}/{dom_size}pt)"
                    ),
                )
            )
    return findings


def standardize(profiles: list[FormatProfile]) -> tuple[dict, list[FormatFinding]]:
    """Cross-document standard + deviations.

    The proposed standard for each kind is the most common dominant style
    across documents — the common denominator. Every document that differs
    gets a named finding.
    """
    votes: dict[str, Counter] = {}
    for profile in profiles:
        for kind, look in profile.dominant.items():
            votes.setdefault(kind, Counter())[look] += 1

    standard = {
        kind: {"font": look[0], "size_pt": look[1], "documents_agreeing": count}
        for kind, counter in sorted(votes.items())
        for look, count in [counter.most_common(1)[0]]
    }

    findings: list[FormatFinding] = []
    for profile in profiles:
        for kind, (font, size) in sorted(profile.dominant.items()):
            std = standard.get(kind)
            if std and (font, size) != (std["font"], std["size_pt"]):
                findings.append(
                    FormatFinding(
                        scope="cross-document",
                        kind=kind,
                        detail=(
                            f"{profile.source}: {kind} is {font}/{size}pt; "
                            f"standard is {std['font']}/{std['size_pt']}pt"
                        ),
                    )
                )
    return standard, findings


def render_report(
    profiles: list[FormatProfile],
    standard: dict,
    findings: list[FormatFinding],
) -> str:
    lines = ["Format standardization report", "=" * 40, ""]
    lines.append("Proposed standard (common denominator across documents):")
    for kind, spec in standard.items():
        lines.append(
            f"  {kind:12} {spec['font']}/{spec['size_pt']}pt "
            f"({spec['documents_agreeing']} of {len(profiles)} documents agree)"
        )
    lines.append("")
    if findings:
        lines.append(f"Deviations ({len(findings)}):")
        for finding in findings:
            lines.append(f"  [{finding.kind}] {finding.detail}")
    else:
        lines.append("No deviations — documents already share one standard.")
    return "\n".join(lines)
